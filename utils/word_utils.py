from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def word_olustur(sinav_baslik, sorular, dosya_yolu="sinav.docx"):
    doc = Document()
    doc.add_heading(sinav_baslik, 0)

    # Sınav bilgileri kısmı
    doc.add_paragraph("Ad Soyad: __________________________________")
    doc.add_paragraph("Öğretmen: __________________________________")
    doc.add_paragraph("Tarih: _____________________________________")
    doc.add_paragraph("Not: _______________________________________\n")

    for idx, soru in enumerate(sorular, start=1):
        # Soru başlığı
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. Soru")
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Zorluk (isteğe bağlı renkli değil, sadece parantez)
        zorluk = soru.get("zorluk")
        if zorluk:
            p.add_run(f" ({zorluk})")
        # Soru metni
        doc.add_paragraph(soru['metin'])
        # Görsel varsa ekle
        if soru['resim'] and os.path.exists(soru['resim']):
            try:
                doc.add_picture(soru['resim'], width=Inches(3.6))
            except Exception as e:
                doc.add_paragraph(f"[Görsel yüklenemedi: {e}]")
        # Seçenekler
        secenekler = ["a", "b", "c", "d"]
        for sec in secenekler:
            metin = soru.get(sec, "")
            doc.add_paragraph(f"{sec.upper()}) {metin}", style='List Bullet')
        doc.add_paragraph("\n")
    doc.save(dosya_yolu)
    return dosya_yolu
